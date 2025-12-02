# -*- coding: utf-8 -*-
"""
AWS Lambda: Build marketing DOCX from S3 template + content
- Auto-detect GitHub-style Markdown tables (no TABLE_START/TABLE_END needed)
- Convert Markdown tables -> real DOCX tables (Light Grid Accent 1)
- Keep headings, bullets, and Label:Value formatting
"""
import uuid 
import boto3
import base64
from boto3.dynamodb.conditions import Key
import json
from io import BytesIO
from botocore.exceptions import ClientError
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from datetime import datetime
import re
import difflib

# ---------- AWS ----------
s3 = boto3.client('s3')
# DynamoDB resource (initialize globally)
dynamodb = boto3.resource('dynamodb')
DOCUMENT_HISTORY_TABLE = "documents-history-table"
users_table = dynamodb.Table("users-table")

# ---------- CONFIG ----------
TABLE_STYLE = "Light Grid Accent 1"

def get_user_id_from_session(session_id: str):
    """
    Fetch user_id from DynamoDB Users table using GSI: session_id-index
    """
    try:
        response = users_table.query(
            IndexName="session_id-index",
            KeyConditionExpression=Key("session_id").eq(session_id)
        )

        items = response.get("Items", [])
        if not items:
            raise Exception(f"No user found for session_id: {session_id}")

        # Assuming your table has a field "id" which represents user_id
        return items[0].get("id")

    except Exception as e:
        print(f"Error fetching user_id for session_id={session_id}: {str(e)}")
        raise


def save_document_history_to_dynamodb(user_id, project_id, document_type, document_name, document_url):
    """
    Inserts a new record into the DocumentHistory DynamoDB table.

    document_type_uuid = "{document_type}#<UUID>"
    """
    try:
        table = dynamodb.Table(DOCUMENT_HISTORY_TABLE)

        document_type_uuid = f"{document_type}#{uuid.uuid4()}"
        created_at = datetime.utcnow().isoformat()

        item = {
            "user_id": user_id,
            "document_type_uuid": document_type_uuid,
            "project_id": project_id,
            "document_type": document_type,
            "document_name": document_name,
            "document_url": document_url,
            "created_at": created_at
        }

        table.put_item(Item=item)

        print(f"✅ Document history saved: {document_type_uuid}")
        return True

    except ClientError as e:
        print(f"❌ DynamoDB error while saving document history: {e}")
        return False
    except Exception as e:
        print(f"❌ Unexpected error while saving document history: {e}")
        return False

# ---------- Helpers: formatting ----------
def apply_base_format(run, size=12, bold=False):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.bold = bold

def format_heading(text: str) -> str:
    return ' '.join(word.capitalize() for word in text.replace("_", " ").split())

# ---------- Text cleanup ----------
def clean_text(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # don't collapse single newlines (tables depend on line structure)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()

def unbreak_paragraphs(text: str) -> str:
    """
    Groups soft-wrapped lines into paragraphs — BUT this should NOT
    run on blocks that contain Markdown tables (we guard for that upstream).
    """
    lines = text.splitlines()
    result = []
    buffer = ""

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if buffer:
                result.append(buffer.strip())
                buffer = ""
            continue

        # bullets are separate
        if stripped.startswith(("-", "*", "•")):
            if buffer:
                result.append(buffer.strip())
                buffer = ""
            result.append(stripped)
            continue

        # "Label: Value" lines are separate
        if ':' in stripped and not stripped.endswith(':'):
            if buffer:
                result.append(buffer.strip())
                buffer = ""
            result.append(stripped)
            continue

        # normal line join
        if buffer and not buffer.endswith(('.', ':', '?', '!', '”', '"')):
            buffer += " " + stripped
        else:
            buffer += ("\n" if buffer else "") + stripped

    if buffer:
        result.append(buffer.strip())

    return "\n\n".join(result)

# ---------- S3 I/O ----------
def read_json_from_s3(bucket, key):
    response = s3.get_object(Bucket=bucket, Key=key)
    content = response['Body'].read().decode('utf-8')
    return json.loads(content)

def read_text_file_from_s3(s3_path):
    """
    Reads UTF-8 text. If it contains a Markdown table, we DO NOT run
    unbreak_paragraphs() (to preserve row lines). Otherwise, we do.
    """
    try:
        s3_path = s3_path.replace("s3://", "")
        bucket, key = s3_path.split("/", 1)

        response = s3.get_object(Bucket=bucket, Key=key)
        raw_text = response['Body'].read().decode('utf-8').strip()
        if not raw_text:
            return "[Content missing]"

        cleaned = clean_text(raw_text)

        # Detect GitHub-style Markdown table: header + separator row like |---|
        lines = cleaned.split("\n")
        has_md_table = False
        sep_re = re.compile(r'^\s*\|?\s*:?-{3,}\s*(\|\s*:?-{3,}\s*)+\|?\s*$')

        for i in range(len(lines) - 1):
            if "|" in lines[i] and sep_re.match(lines[i + 1] or ""):
                has_md_table = True
                break

        # Only unbreak when there's no table (tables rely on exact line structure)
        if has_md_table:
            return cleaned
        else:
            return unbreak_paragraphs(cleaned)

    except ClientError as e:
        if e.response['Error']['Code'] == 'NoSuchKey':
            return "[Content missing]"
        raise e

# ---------- Exact heading matching (skip duplicate H1s inside content) ----------
def _normalize_heading(s: str) -> str:
    s = re.sub(r'\s+', ' ', s).strip().lower()
    s = re.sub(r'[^a-z0-9 ]+', '', s)
    return s

def exact_match(a: str, b: str) -> bool:
    """
    Exact matching for headings - only skip if they're identical
    after normalization (spaces, case, punctuation removed)
    """
    na, nb = _normalize_heading(a), _normalize_heading(b)
    return na == nb

# ---------- Markdown table -> real DOCX table ----------
def markdown_table_to_docx(md_table: str, doc: Document):
    lines = [line.strip() for line in md_table.strip().split("\n") if line.strip()]
    if len(lines) < 2:
        return

    # First non-empty line is header; skip the next separator row(s)
    header = [h.strip() for h in lines[0].split("|") if h.strip()]
    if not header:
        return

    # Build rows (skip separator lines)
    rows = []
    sep_re = re.compile(r'^\s*\|?\s*:?-{3,}\s*(\|\s*:?-{3,}\s*)+\|?\s*$')
    for line in lines[1:]:
        if sep_re.match(line):
            continue
        if "|" in line:
            row = [c.strip() for c in line.split("|") if c.strip()]
            if row:
                rows.append(row)

    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    try:
        table.style = TABLE_STYLE
    except Exception:
        pass
    table.autofit = True

    # header
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                apply_base_format(run, size=11, bold=True)

    # body
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = txt
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    apply_base_format(run, size=10, bold=False)

# ---------- Content renderer ----------
def add_formatted_paragraphs(document: Document, text: str, template_h1: str = ""):
    """
    - Auto-detects GitHub-style Markdown tables (header + --- separator)
    - Converts them into real DOCX tables
    - Recognizes markdown headings (#, ##, ###)
    - Skips H1 if exact-matching template_h1 (avoid dupes)
    - Bullets: -, *, •
    - "Label:" or "Label: Value" → bold label
    """
    text = clean_text(text)
    lines = text.split("\n")

    i = 0
    sep_re = re.compile(r'^\s*\|?\s*:?-{3,}\s*(\|\s*:?-{3,}\s*)+\|?\s*$')

    def flush_table_from(start_idx):
        """
        Given index of header row, consume:
          header line
          one or more separator lines
          subsequent lines with pipes (rows)
        Return (end_idx_exclusive, md_table_string)
        """
        header_line = lines[start_idx].strip()
        j = start_idx + 1

        # must have at least one separator line
        if j >= len(lines) or not sep_re.match(lines[j].strip()):
            return start_idx, None

        buf = [header_line]
        # include all consecutive separator lines (rare but allowed)
        while j < len(lines) and sep_re.match(lines[j].strip()):
            buf.append(lines[j].strip())
            j += 1

        # include table rows: lines that contain at least one pipe
        while j < len(lines) and '|' in lines[j]:
            buf.append(lines[j].strip())
            j += 1

        return j, "\n".join(buf)

    while i < len(lines):
        line = lines[i].strip()

        # try table detection first (header + separator)
        if "|" in line and (i + 1) < len(lines) and sep_re.match(lines[i + 1].strip()):
            end_idx, md_table = flush_table_from(i)
            if md_table:
                markdown_table_to_docx(md_table, document)
                document.add_paragraph()
                i = end_idx
                continue  # next line after table

        # blank line -> spacing
        if not line:
            document.add_paragraph()
            i += 1
            continue

        # H3
        if line.startswith("### "):
            p = document.add_paragraph()
            run = p.add_run(line[4:].strip())
            apply_base_format(run, size=14, bold=True)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            p = document.add_paragraph()
            run = p.add_run(line[3:].strip())
            apply_base_format(run, size=16, bold=True)
            i += 1
            continue

        # H1 (skip if identical to template heading)
        if line.startswith("# "):
            h1_text = line[2:].strip()
            if not (template_h1 and exact_match(h1_text, template_h1)):
                p = document.add_paragraph()
                run = p.add_run(h1_text)
                apply_base_format(run, size=17, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # bullets
        if line.startswith(("-", "*", "•")):
            content = line[1:].strip()
            p = document.add_paragraph(style='List Bullet')
            if ':' in content:
                # Handle both "Label:" and "Label: Value" formats
                if content.endswith(':'):
                    # Case: "Revenue Streams:" - make entire content bold
                    run = p.add_run(content)
                    apply_base_format(run, bold=True)
                else:
                    # Case: "Label: Value" - make only label part bold
                    label, _, rest = content.partition(':')
                    run_b = p.add_run(label.strip() + ": ")
                    apply_base_format(run_b, bold=True)
                    run_n = p.add_run(rest.strip())
                    apply_base_format(run_n)
            else:
                run = p.add_run(content)
                apply_base_format(run)
            i += 1
            continue

        # Label: Value OR Label: (standalone labels ending with colon)
        if ':' in line:
            if line.endswith(':'):
                # Case: "Revenue Streams:" - entire line bold
                p = document.add_paragraph()
                run = p.add_run(line)
                apply_base_format(run, bold=True)
            else:
                # Case: "Label: Value" - only label part bold
                label, value = line.split(':', 1)   # split only once
                label = label.strip()
                value = value.strip()

                p = document.add_paragraph()

                # Bold label
                run_b = p.add_run(label + ": ")
                apply_base_format(run_b, bold=True)

                # Normal value (explicitly un-bolded)
                if value:
                    run_n = p.add_run(value)
                    apply_base_format(run_n, bold=False)

            i += 1
            continue        

        # default paragraph
        p = document.add_paragraph()
        run = p.add_run(line)
        apply_base_format(run)
        i += 1

# ---------- Status updater ----------
def update_status_to_false(bucket_name, object_key):
    try:
        response = s3.get_object(Bucket=bucket_name, Key=object_key)
        content = response['Body'].read().decode('utf-8')
        data = json.loads(content)

        for tier in data.values():
            if isinstance(tier, dict) and 'status' in tier:
                tier['status'] = False

        updated_content = json.dumps(data, indent=2)
        s3.put_object(Bucket=bucket_name, Key=object_key, Body=updated_content.encode('utf-8'))

        print("✅ Status keys updated successfully.")

    except ClientError as e:
        print(f"❌ AWS Error: {e}")
    except json.JSONDecodeError as e:
        print(f"❌ JSON Error: {e}")
    except Exception as e:
        print(f"❌ Unexpected Error: {e}")

# ---------- Lambda handler ----------
def lambda_handler(event, context):
    
    project_id = event.get("project_id", "")
    session_id = event.get("session_id", "")
    user_id = get_user_id_from_session(session_id)
    document_type = event.get("document_type", "")

    template_bucket = 'cammi-devprod'
    # object_key = f'flow/{user_id}/{document_type}/execution_plan.json'
    template_key = f'flow/{document_type}/marketing_document_template.json'
    output_bucket = 'cammi-devprod'

    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    output_key = f'project/{project_id}/{document_type}/marketing_strategy_document/{document_type}_{timestamp}.docx'

    # Load template JSON
    template = read_json_from_s3(template_bucket, template_key)

    # Build DOCX
    document = Document()

    # Default font
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # Render sections
    for section in template:
        for subsection in section.get("sections", []):
            subheading = format_heading(subsection.get("subheading", ""))
            s3_path = subsection.get("s3_path", "")
            if not s3_path.startswith("s3://"):
                s3_path = f"s3://cammi-devprod/{project_id}/{document_type}/{s3_path}"
            content_text = read_text_file_from_s3(s3_path)

            # spacing before heading
            document.add_paragraph()

            # Heading
            p = document.add_paragraph(style='Heading 1')
            run = p.add_run(subheading)
            apply_base_format(run, size=17, bold=True)
            run.font.color.rgb = RGBColor(0, 0, 0)

            # spacing after heading
            document.add_paragraph()

            # Content (auto tables + rich formatting)
            add_formatted_paragraphs(document, content_text, template_h1=subheading)

    # Save to buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Upload to S3
    s3.put_object(
        Bucket=output_bucket,
        Key=output_key,
        Body=buffer.getvalue(),
        ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    # Build document info
    document_name = output_key.split('/')[-1]
    document_url = f"s3://{output_bucket}/{output_key}"

    # Save document history to DynamoDB
    save_document_history_to_dynamodb(
        user_id=user_id,
        project_id=project_id,
        document_type=document_type,
        document_name=document_name,
        document_url=document_url
    )
    
	# Folder path for the user
    FOLDER_PREFIX = f"project/{project_id}/{document_type}/marketing_strategy_document/"
	# List objects in S3 folder
    response = s3.list_objects_v2(Bucket=output_bucket, Prefix=FOLDER_PREFIX)
	# Filter only .docx files
    docx_files = [
		obj for obj in response.get("Contents", [])
		if obj["Key"].endswith(".docx")
    ]
    if not docx_files:
	    return {
			"statusCode": 404,
			"body": json.dumps({"error": "No .docx files found in the folder"}),
			"headers": CORS_HEADERS
	    }
	# Get latest file
    latest_file = max(docx_files, key=lambda x: x["LastModified"])
    latest_key = latest_file["Key"]
	# Read and encode file
    s3_response = s3.get_object(Bucket=output_bucket, Key=latest_key)
    file_data = s3_response["Body"].read()
    encoded_data = base64.b64encode(file_data).decode("utf-8")



    # Reset execution plan statuses
    # update_status_to_false(bucket_name='cammi', object_key=object_key)

    return {
		"statusCode": 200,
		"body": json.dumps({
			"message": "Latest .docx file fetched successfully",
			"fileName": latest_key.split("/")[-1],
			"docxBase64": encoded_data
		})
    }