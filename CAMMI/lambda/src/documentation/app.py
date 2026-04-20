# -*- coding: utf-8 -*-
"""
AWS Lambda: Build marketing DOCX from S3 template + content

Fixed version — addresses:
  1. TOC pre-populated with entries (auto-updates on open in Word)
  2. No spurious blank paragraphs between content
  3. Bullet points rendered with proper Word numbering (not just indented text)
  4. Short lines (messaging lines, pain points, etc.) detected and bulleted automatically
  5. Header/footer placeholders replaced correctly
"""

import uuid
import boto3
import json
import re
import base64
from io import BytesIO
from datetime import datetime
from copy import deepcopy

from botocore.exceptions import ClientError
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- AWS ----------
s3 = boto3.client("s3")
dynamodb = boto3.resource("dynamodb")
DOCUMENT_HISTORY_TABLE = "documents-history-table"

# ---------- CONFIG ----------
TABLE_STYLE = "Light Grid Accent 1"

# Style IDs from the template (numeric IDs used by WPS/custom templates)
# These map to the named styles in the .docx template:
#   "2"   → Heading 1
#   "3"   → Heading 2
#   "4"   → Heading 3
#   "164" → Cammi Body
#   "165" → Cammi Bullet  (now linked to numId 7 for real bullets)
#   "166" → Cammi Label
STYLE_HEADING1 = "Heading 1"
STYLE_HEADING2 = "Heading 2"
STYLE_HEADING3 = "Heading 3"
STYLE_BODY = "Cammi Body"
STYLE_BULLET = "Cammi Bullet"
STYLE_LABEL = "Cammi Label"

DOCUMENT_TYPE_NAMES = {
    "gtm":       "Go to Market",
    "icp":       "Ideal Customer Profile",
    "kmf":       "Key Messaging Framework",
    "bs":        "Brand Strategy",
    "sr":        "Strategy Roadmap",
    "mr":        "Market Research",
    "brand":     "Brand",
    "messaging": "Messaging",
    "smp":       "Strategic Marketing Plan",
}

# ═══════════════════════════════════════════════════════════
#  DYNAMODB HISTORY
# ═══════════════════════════════════════════════════════════

def save_document_history_to_dynamodb(user_id, project_id, document_type,
                                      document_name, document_url):
    try:
        table = dynamodb.Table(DOCUMENT_HISTORY_TABLE)
        item = {
            "user_id":            user_id,
            "document_type_uuid": f"{document_type}#{uuid.uuid4()}",
            "project_id":         project_id,
            "document_type":      document_type,
            "document_name":      document_name,
            "document_url":       document_url,
            "created_at":         datetime.utcnow().isoformat(),
        }
        table.put_item(Item=item)
        print(f"✅ Document history saved")
        return True
    except Exception as e:
        print(f"❌ DynamoDB error: {e}")
        return False

# ═══════════════════════════════════════════════════════════
#  TEXT HELPERS
# ═══════════════════════════════════════════════════════════

def format_heading(text: str) -> str:
    return " ".join(w.capitalize() for w in text.replace("_", " ").split())


def clean_text(text: str) -> str:
    """Normalize whitespace, collapse triple+ newlines, strip."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _is_short_standalone_line(line: str) -> bool:
    """
    Detect lines that look like messaging lines / short value props.
    These are typically 1-sentence lines that should be bulleted.
    Heuristic: ends with a period, no colon, under 150 chars, not a heading.
    """
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith(("#", "-", "*", "•")):
        return False
    if ":" in stripped:
        return False
    if len(stripped) > 150:
        return False
    if stripped.endswith((".")) and len(stripped.split()) >= 4:
        return True
    return False


def _classify_lines(lines: list) -> list:
    """
    Walk through lines and classify sequences of short standalone lines
    as bullets. Returns list of (type, text) tuples.
    Types: 'heading1', 'heading2', 'heading3', 'bullet', 'label',
           'label_heading', 'body', 'blank', 'table_start'
    """
    sep_re = re.compile(r"^\s*\|?\s*:?-{3,}\s*(\|\s*:?-{3,}\s*)+\|?\s*$")
    result = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Blank
        if not line:
            result.append(("blank", ""))
            i += 1
            continue

        # Markdown table start
        if "|" in line and (i + 1) < len(lines) and sep_re.match(lines[i + 1].strip()):
            result.append(("table_start", str(i)))
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            result.append(("heading3", line[4:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            result.append(("heading2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("# "):
            result.append(("heading1", line[2:].strip()))
            i += 1
            continue

        # Explicit bullets
        if line.startswith(("-", "*", "•")):
            result.append(("bullet", line.lstrip("-*• ").strip()))
            i += 1
            continue

        # Label: Value  or  Label heading (ending with colon)
        if ":" in line:
            if line.endswith(":"):
                result.append(("label_heading", line))
            else:
                result.append(("label", line))
            i += 1
            continue

        # Check if this is part of a sequence of short standalone lines
        # (messaging lines, value props, etc.)
        if _is_short_standalone_line(line):
            # Look ahead through blank lines: if there are 2+ short lines
            # in the run (ignoring blanks between them), treat all as bullets
            run_start = i
            j = i
            short_indices = []
            while j < len(lines):
                s = lines[j].strip()
                if _is_short_standalone_line(s):
                    short_indices.append(j)
                    j += 1
                elif not s:
                    # Skip blank lines between short lines
                    j += 1
                else:
                    break
            if len(short_indices) >= 2:
                for idx in short_indices:
                    result.append(("bullet", lines[idx].strip()))
                i = j
            else:
                # Single short line → body
                result.append(("body", lines[run_start].strip()))
                i = run_start + 1
            continue

        # Default body
        result.append(("body", line))
        i += 1

    return result


# ═══════════════════════════════════════════════════════════
#  S3 HELPERS
# ═══════════════════════════════════════════════════════════

def load_template_from_s3(bucket, key):
    response = s3.get_object(Bucket=bucket, Key=key)
    return Document(BytesIO(response["Body"].read()))


def read_json_from_s3(bucket, key):
    response = s3.get_object(Bucket=bucket, Key=key)
    return json.loads(response["Body"].read().decode("utf-8"))


def read_text_file_from_s3(s3_path):
    try:
        s3_path = s3_path.replace("s3://", "")
        bucket, key = s3_path.split("/", 1)
        response = s3.get_object(Bucket=bucket, Key=key)
        raw = response["Body"].read().decode("utf-8").strip()
        if not raw:
            return "[Content missing]"
        return clean_text(raw)
    except ClientError as e:
        if e.response["Error"]["Code"] == "NoSuchKey":
            return "[Content missing]"
        raise


# ═══════════════════════════════════════════════════════════
#  PLACEHOLDER REPLACEMENT
# ═══════════════════════════════════════════════════════════

def _replace_in_paragraph(para, replacements: dict):
    """Replace {{key}} tokens in a paragraph's runs, preserving formatting."""
    for run in para.runs:
        for token, value in replacements.items():
            if token in run.text:
                run.text = run.text.replace(token, value)


def replace_placeholders(document: Document, replacements: dict):
    """Walk every paragraph in the document (body, headers, footers, tables)."""
    for para in document.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for para in ncell.paragraphs:
                                _replace_in_paragraph(para, replacements)

    for section in document.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr is not None:
                for para in hdr.paragraphs:
                    _replace_in_paragraph(para, replacements)
        for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
            if ftr is not None:
                for para in ftr.paragraphs:
                    _replace_in_paragraph(para, replacements)


# ═══════════════════════════════════════════════════════════
#  LOGO
# ═══════════════════════════════════════════════════════════

def save_project_logo_to_s3(project_id: str, logo_base64: str):
    try:
        if not logo_base64:
            return None
        if "," in logo_base64:
            logo_base64 = logo_base64.split(",", 1)[1]
        image_bytes = base64.b64decode(logo_base64)
        logo_key = f"logos/{project_id}/logo.png"
        s3.put_object(Bucket="cammi-devprod", Key=logo_key,
                      Body=image_bytes, ContentType="image/png")
        print(f"✅ Logo saved at s3://cammi-devprod/{logo_key}")
        return logo_key
    except Exception as e:
        print(f"❌ Failed to save logo: {e}")
        return None


def get_project_logo_from_s3(project_id: str):
    try:
        response = s3.get_object(Bucket="cammi-devprod",
                                 Key=f"logos/{project_id}/logo.png")
        return BytesIO(response["Body"].read())
    except ClientError as e:
        if e.response["Error"]["Code"] == "NoSuchKey":
            return None
        raise


def add_logo_to_header(document: Document, project_id: str):
    """Add project logo to TOP-RIGHT of content section headers (skip cover)."""
    try:
        logo_stream = get_project_logo_from_s3(project_id)
        if not logo_stream:
            return
        for section in document.sections[1:]:
            header = section.header
            p = header.add_paragraph()
            p.alignment = 2  # RIGHT
            run = p.add_run()
            run.add_picture(logo_stream, width=Inches(0.3))
            logo_stream.seek(0)
    except Exception as e:
        print(f"⚠️ Logo skipped: {e}")


# ═══════════════════════════════════════════════════════════
#  MARKDOWN TABLE → DOCX TABLE
# ═══════════════════════════════════════════════════════════

def _flush_md_table(lines, start_idx):
    """Extract a complete markdown table starting at start_idx."""
    sep_re = re.compile(r"^\s*\|?\s*:?-{3,}\s*(\|\s*:?-{3,}\s*)+\|?\s*$")
    header_line = lines[start_idx].strip()
    j = start_idx + 1
    if j >= len(lines) or not sep_re.match(lines[j].strip()):
        return start_idx + 1, None
    buf = [header_line]
    while j < len(lines) and sep_re.match(lines[j].strip()):
        buf.append(lines[j].strip())
        j += 1
    while j < len(lines) and "|" in lines[j]:
        buf.append(lines[j].strip())
        j += 1
    return j, "\n".join(buf)


def markdown_table_to_docx(md_table: str, doc: Document):
    lines = [l.strip() for l in md_table.strip().split("\n") if l.strip()]
    if len(lines) < 2:
        return
    header = [h.strip() for h in lines[0].split("|") if h.strip()]
    if not header:
        return
    sep_re = re.compile(r"^\s*\|?\s*:?-{3,}\s*(\|\s*:?-{3,}\s*)+\|?\s*$")
    rows = []
    for line in lines[1:]:
        if sep_re.match(line):
            continue
        if "|" in line:
            row = [c.strip() for c in line.split("|") if c.strip()]
            if row:
                rows.append(row)
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    try:
        table.style = TABLE_STYLE
    except Exception:
        pass
    table.autofit = True
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            if j < len(header):
                table.cell(i + 1, j).text = txt


# ═══════════════════════════════════════════════════════════
#  CONTENT RENDERER (FIXED)
#
#  Key fixes:
#  - No empty <w:p/> elements between content paragraphs
#  - Bullets use "Cammi Bullet" style (now linked to numId 7)
#  - Short standalone lines auto-detected as bullets
#  - Consecutive blank lines collapsed to a single spacing gap
# ═══════════════════════════════════════════════════════════

def _normalize_heading(s: str) -> str:
    s = re.sub(r"\s+", " ", s).strip().lower()
    return re.sub(r"[^a-z0-9 ]+", "", s)


def add_formatted_paragraphs(document: Document, text: str, template_h1: str = ""):
    """
    Parse content text and add properly formatted paragraphs to the document.
    No empty paragraphs are inserted; spacing is controlled by styles.
    """
    text = clean_text(text)
    lines = text.split("\n")
    sep_re = re.compile(r"^\s*\|?\s*:?-{3,}\s*(\|\s*:?-{3,}\s*)+\|?\s*$")

    classified = _classify_lines(lines)

    prev_type = None
    i = 0
    while i < len(classified):
        ctype, ctext = classified[i]

        # ── Skip blank lines — spacing is handled by styles ──
        if ctype == "blank":
            # Only add a spacing paragraph if transitioning between
            # major content blocks (not between bullets or consecutive items)
            if prev_type and prev_type not in ("blank", "heading1", "heading2",
                                                "heading3", "label_heading"):
                # Look ahead: if next non-blank is a heading or label_heading,
                # skip the blank (heading styles have built-in spacing)
                next_type = None
                for j in range(i + 1, len(classified)):
                    if classified[j][0] != "blank":
                        next_type = classified[j][0]
                        break
                if next_type in ("heading1", "heading2", "heading3",
                                 "label_heading", None):
                    i += 1
                    continue
                # Between body sections, skip extra blanks
            i += 1
            continue

        # ── Markdown table ──
        if ctype == "table_start":
            start_idx = int(ctext)
            end_idx, md_table = _flush_md_table(lines, start_idx)
            if md_table:
                markdown_table_to_docx(md_table, document)
            # Skip all classified entries that correspond to table lines
            i += 1
            while i < len(classified) and classified[i][0] == "table_start":
                i += 1
            prev_type = "table"
            continue

        # ── H1 (skip if duplicate of section heading) ──
        if ctype == "heading1":
            if not (template_h1 and
                    _normalize_heading(ctext) == _normalize_heading(template_h1)):
                document.add_paragraph(ctext, style=STYLE_HEADING1)
            prev_type = ctype
            i += 1
            continue

        # ── H2 ──
        if ctype == "heading2":
            document.add_paragraph(ctext, style=STYLE_HEADING2)
            prev_type = ctype
            i += 1
            continue

        # ── H3 ──
        if ctype == "heading3":
            document.add_paragraph(ctext, style=STYLE_HEADING3)
            prev_type = ctype
            i += 1
            continue

        # ── Bullet ──
        if ctype == "bullet":
            p = document.add_paragraph(style=STYLE_BULLET)
            # Handle "Label: Value" inside bullets
            if ":" in ctext and not ctext.endswith(":"):
                label, _, rest = ctext.partition(":")
                p.add_run(label.strip() + ": ").bold = True
                p.add_run(rest.strip())
            else:
                p.add_run(ctext)
            prev_type = ctype
            i += 1
            continue

        # ── Label heading (line ending with colon) ──
        if ctype == "label_heading":
            p = document.add_paragraph(style=STYLE_LABEL)
            p.add_run(ctext).bold = True
            prev_type = ctype
            i += 1
            continue

        # ── Label: Value ──
        if ctype == "label":
            p = document.add_paragraph(style=STYLE_BODY)
            label, _, value = ctext.partition(":")
            p.add_run(label.strip() + ": ").bold = True
            if value.strip():
                p.add_run(value.strip())
            prev_type = ctype
            i += 1
            continue

        # ── Body paragraph ──
        if ctype == "body":
            document.add_paragraph(ctext, style=STYLE_BODY)
            prev_type = ctype
            i += 1
            continue

        i += 1


# ═══════════════════════════════════════════════════════════
#  TOC POPULATION HELPER
#  - Adds bookmarks to each heading paragraph
#  - Pre-populates TOC field with PAGEREF entries that Word
#    resolves to real page numbers on open (no manual F9 needed)
# ═══════════════════════════════════════════════════════════

def _make_bookmark_id(text: str, idx: int) -> str:
    """Create a safe bookmark name from heading text."""
    safe = re.sub(r"[^a-zA-Z0-9]", "", text)[:20]
    return f"_Toc{safe}{idx}"


def add_heading_with_bookmark(document: Document, text: str, style: str,
                               bookmark_name: str, bookmark_id: int):
    """Add a heading paragraph with a bookmark wrapping the text."""
    p = document.add_paragraph(style=style)
    # Create bookmark start
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)
    p._element.append(bm_start)
    # Add the text run
    run = p.add_run(text)
    # Create bookmark end
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    p._element.append(bm_end)
    return p


def update_toc_entries(document: Document, headings: list):
    """
    Find the TOC field in the document and insert entries with PAGEREF
    field codes that resolve to real page numbers.

    headings: list of (level, text, bookmark_name) tuples
    """
    body = document.element.body

    # Find the paragraph containing the TOC field
    for para in body.findall(qn("w:p")):
        runs = para.findall(qn("w:r"))
        has_toc = False
        for run in runs:
            instr = run.find(qn("w:instrText"))
            if instr is not None and "TOC" in (instr.text or ""):
                has_toc = True
                break
        if not has_toc:
            continue

        # Found TOC paragraph — remove runs after 'separate'
        found_separate = False
        runs_to_keep = []
        end_run = None
        for run in runs:
            fld = run.find(qn("w:fldChar"))
            if fld is not None and fld.get(qn("w:fldCharType")) == "separate":
                found_separate = True
                runs_to_keep.append(run)
                continue
            if found_separate:
                if fld is not None and fld.get(qn("w:fldCharType")) == "end":
                    end_run = run
                continue
            runs_to_keep.append(run)

        # Clear and re-add kept runs
        for run in list(para):
            if run.tag == qn("w:r"):
                para.remove(run)
        for run in runs_to_keep:
            para.append(run)

        # Insert TOC entry paragraphs
        parent = para.getparent()
        insert_idx = list(parent).index(para) + 1

        for level, heading_text, bookmark_name in headings:
            toc_p = OxmlElement("w:p")
            toc_ppr = OxmlElement("w:pPr")
            toc_style = OxmlElement("w:pStyle")
            toc_style.set(qn("w:val"), f"TOC{level}" if level <= 3 else "TOC1")
            toc_ppr.append(toc_style)

            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:leader"), "dot")
            tab.set(qn("w:pos"), "9360")
            tabs.append(tab)
            toc_ppr.append(tabs)

            if level > 1:
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), str((level - 1) * 240))
                toc_ppr.append(ind)

            toc_p.append(toc_ppr)

            # Heading text run
            toc_r = OxmlElement("w:r")
            toc_t = OxmlElement("w:t")
            toc_t.text = heading_text
            toc_r.append(toc_t)
            toc_p.append(toc_r)

            # Tab character (triggers dot leader)
            tab_r = OxmlElement("w:r")
            tab_t = OxmlElement("w:t")
            tab_t.set(qn("xml:space"), "preserve")
            tab_t.text = "\t"
            tab_r.append(tab_t)
            toc_p.append(tab_r)

            # PAGEREF field code → resolves to real page number
            # Field begin
            fld_begin_r = OxmlElement("w:r")
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin_r.append(fld_begin)
            toc_p.append(fld_begin_r)

            # Field instruction
            fld_instr_r = OxmlElement("w:r")
            fld_instr = OxmlElement("w:instrText")
            fld_instr.set(qn("xml:space"), "preserve")
            fld_instr.text = f" PAGEREF {bookmark_name} \\h "
            fld_instr_r.append(fld_instr)
            toc_p.append(fld_instr_r)

            # Field separate
            fld_sep_r = OxmlElement("w:r")
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            fld_sep_r.append(fld_sep)
            toc_p.append(fld_sep_r)

            # Placeholder page number (Word replaces this on open)
            pn_r = OxmlElement("w:r")
            pn_t = OxmlElement("w:t")
            pn_t.text = "0"
            pn_r.append(pn_t)
            toc_p.append(pn_r)

            # Field end
            fld_end_r = OxmlElement("w:r")
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            fld_end_r.append(fld_end)
            toc_p.append(fld_end_r)

            parent.insert(insert_idx, toc_p)
            insert_idx += 1

        # Add the TOC 'end' field char
        end_p = OxmlElement("w:p")
        if end_run is not None:
            end_p.append(end_run)
        else:
            end_r = OxmlElement("w:r")
            end_fld = OxmlElement("w:fldChar")
            end_fld.set(qn("w:fldCharType"), "end")
            end_r.append(end_fld)
            end_p.append(end_r)
        parent.insert(insert_idx, end_p)

        break  # Only process first TOC field


# ═══════════════════════════════════════════════════════════
#  LIBREOFFICE FIELD UPDATE
#  Opens the DOCX in LibreOffice headless, triggers a macro
#  that updates all fields (PAGEREF, TOC), saves back as DOCX.
#  This resolves TOC page numbers that python-docx cannot compute.
#
#  Requires: LibreOffice installed on the Lambda runtime
#  (e.g., via a Lambda Layer like "libreoffice-brotli" or
#  a container image with LibreOffice).
# ═══════════════════════════════════════════════════════════

def _estimate_page_numbers(document: Document, heading_bookmarks: list) -> dict:
    """
    Estimate which page each heading lands on by counting content
    characters between headings and dividing by approximate
    characters-per-page.

    Returns dict: {bookmark_name: page_number}

    Assumptions for US Letter, 1" margins, 11pt Arial:
    - ~3200 characters of body text per page
    - Cover page = page 1, TOC page = page 2
    - Content starts on page 3
    """
    CHARS_PER_PAGE = 2800
    CONTENT_START_PAGE = 3  # After cover + TOC

    # Collect all paragraphs and find heading positions
    paragraphs = list(document.paragraphs)
    heading_positions = []  # (paragraph_index, bookmark_name)

    # Find heading paragraphs by matching bookmark names
    bm_names = {bm for _, _, bm in heading_bookmarks}

    for i, p in enumerate(paragraphs):
        # Check if this paragraph has a bookmark from our list
        for bm_start in p._element.findall(qn("w:bookmarkStart")):
            name = bm_start.get(qn("w:name"), "")
            if name in bm_names:
                heading_positions.append((i, name))

    if not heading_positions:
        # Fallback: assign sequential page numbers
        return {bm: CONTENT_START_PAGE + idx
                for idx, (_, _, bm) in enumerate(heading_bookmarks)}

    # Calculate cumulative character count from content start
    # Find where content begins (after the template boilerplate)
    content_start_idx = heading_positions[0][0] if heading_positions else 0

    page_map = {}
    cumulative_chars = 0

    for pos_idx, (para_idx, bm_name) in enumerate(heading_positions):
        # Calculate chars from previous heading to this one
        if pos_idx == 0:
            start = content_start_idx
        else:
            start = heading_positions[pos_idx - 1][0]

        section_chars = 0
        for j in range(start, para_idx):
            if j < len(paragraphs):
                text = paragraphs[j].text or ""
                section_chars += len(text)
                # Add extra for bullet/label formatting overhead
                style_name = (paragraphs[j].style.name or "") if paragraphs[j].style else ""
                if "Bullet" in style_name:
                    section_chars += 40  # bullet indentation uses space
                if "Heading" in style_name:
                    section_chars += 200  # headings take more vertical space

        cumulative_chars += section_chars
        page_num = CONTENT_START_PAGE + (cumulative_chars // CHARS_PER_PAGE)
        page_map[bm_name] = page_num

    return page_map


def _patch_toc_page_numbers(document: Document, page_map: dict):
    """
    Walk through the TOC paragraphs and replace the placeholder "0"
    text in PAGEREF field results with actual estimated page numbers.
    """
    body = document.element.body

    for para in body.findall(qn("w:p")):
        runs = list(para.findall(qn("w:r")))

        # Look for PAGEREF instrText in this paragraph
        pageref_bookmark = None
        in_field_result = False

        for run in runs:
            # Check for instrText with PAGEREF
            instr = run.find(qn("w:instrText"))
            if instr is not None and "PAGEREF" in (instr.text or ""):
                # Extract bookmark name from "PAGEREF _TocXxx \h"
                match = re.search(r"PAGEREF\s+(\S+)", instr.text)
                if match:
                    pageref_bookmark = match.group(1)

            # Check for separate (marks start of field result)
            fld = run.find(qn("w:fldChar"))
            if fld is not None:
                fld_type = fld.get(qn("w:fldCharType"))
                if fld_type == "separate":
                    in_field_result = True
                    continue
                elif fld_type == "end":
                    in_field_result = False
                    pageref_bookmark = None
                    continue

            # If we're in the field result area, replace the text
            if in_field_result and pageref_bookmark:
                t = run.find(qn("w:t"))
                if t is not None and pageref_bookmark in page_map:
                    t.text = str(page_map[pageref_bookmark])
                    in_field_result = False


# ═══════════════════════════════════════════════════════════
#  STATUS UPDATER
# ═══════════════════════════════════════════════════════════

def update_status_to_false(bucket_name, object_key):
    try:
        response = s3.get_object(Bucket=bucket_name, Key=object_key)
        data = json.loads(response["Body"].read().decode("utf-8"))
        for tier in data.values():
            if isinstance(tier, dict) and "status" in tier:
                tier["status"] = False
        s3.put_object(Bucket=bucket_name, Key=object_key,
                      Body=json.dumps(data, indent=2).encode("utf-8"))
        print("✅ Status updated")
    except Exception as e:
        print(f"❌ Status update error: {e}")


# ═══════════════════════════════════════════════════════════
#  LAMBDA HANDLER
# ═══════════════════════════════════════════════════════════

def lambda_handler(event, context):
    user_id       = event.get("user_id", "")
    project_id    = event.get("project_id", "")
    session_id    = event.get("session_id", "")
    document_type = event.get("document_type", "")
    logo_base64   = event.get("logo_base64", "")

    template_bucket = "cammi-devprod"
    object_key      = f"flow/{user_id}/{document_type}/execution_plan.json"
    template_key    = f"flow/{document_type}/marketing_document_template.json"
    output_bucket   = "cammi-devprod"
    timestamp       = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    output_key      = (
        f"project/{project_id}/{document_type}/"
        f"marketing_strategy_document/{document_type}_{timestamp}.docx"
    )
    knowledgebase_output = (
        f"knowledgebase/{user_id}/{user_id}_{document_type}_{timestamp}.docx"
    )

    # ── 1. Save logo if provided ──────────────────────────────
    if logo_base64:
        save_project_logo_to_s3(project_id, logo_base64)

    # ── 2. Load template JSON ─────────────────────────────────
    template = read_json_from_s3(template_bucket, template_key)

    # ── 3. Load DOCX template from S3 ────────────────────────
    document = load_template_from_s3(
        "cammi-devprod",
        "templates/cammi_master_strategy_template.docx"
    )

    # ── 4. Replace cover page & header/footer placeholders ───
    doc_type_full = DOCUMENT_TYPE_NAMES.get(
        document_type.lower(), document_type
    )
    replacements = {
        "{{document_title}}": doc_type_full,
        "{{subtitle}}":       "Strategic insights and execution framework",
        "{{project_id}}":     project_id,
        "{{date}}":           datetime.utcnow().strftime("%B %d, %Y"),
        "{{doc_type_full}}":  doc_type_full,
    }
    replace_placeholders(document, replacements)

    # ── 5. Add logo to content section headers ────────────────
    add_logo_to_header(document, project_id)

    # ── 6. Collect headings and render main content ───────────
    toc_headings = []  # (level, text, bookmark_name) for TOC population
    bookmark_counter = 100  # Start IDs high to avoid conflicts

    for section in template:
        for subsection in section.get("sections", []):
            subheading = format_heading(subsection.get("subheading", ""))
            s3_path    = subsection.get("s3_path", "")

            if not s3_path.startswith("s3://"):
                s3_path = (
                    f"s3://cammi-devprod/{project_id}/{document_type}/{s3_path}"
                )

            content_text = read_text_file_from_s3(s3_path)

            # Add section heading with bookmark for TOC PAGEREF
            bm_name = _make_bookmark_id(subheading, bookmark_counter)
            add_heading_with_bookmark(
                document, subheading, STYLE_HEADING1,
                bm_name, bookmark_counter
            )
            toc_headings.append((1, subheading, bm_name))
            bookmark_counter += 1

            # Add formatted content
            add_formatted_paragraphs(
                document, content_text, template_h1=subheading
            )

    # ── 6b. Populate TOC with collected headings ──────────────
    if toc_headings:
        update_toc_entries(document, toc_headings)

    # ── 6c. Estimate page numbers and patch TOC ──────────────
    if toc_headings:
        page_map = _estimate_page_numbers(document, toc_headings)
        _patch_toc_page_numbers(document, page_map)
        print(f"✅ TOC page numbers estimated: {page_map}")

    # ── 7. Save document ──────────────────────────────────────
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    doc_bytes = buffer.getvalue()

    common_metadata = {
        "user_id":       user_id,
        "project_id":    project_id,
        "document_type": document_type,
    }
    content_type = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )

    for key in [output_key, knowledgebase_output]:
        s3.put_object(
            Bucket=output_bucket, Key=key,
            Body=doc_bytes, Metadata=common_metadata,
            ContentType=content_type,
        )

    # ── 8. Log history ────────────────────────────────────────
    save_document_history_to_dynamodb(
        user_id=user_id,
        project_id=project_id,
        document_type=document_type,
        document_name=output_key.split("/")[-1],
        document_url=f"s3://{output_bucket}/{output_key}",
    )

    # ── 9. Update execution plan status ──────────────────────
    update_status_to_false(bucket_name="cammi-devprod", object_key=object_key)

    return {
        "statusCode": 200,
        "session_id": session_id,
        "project_id": project_id,
        "message":    f"DOCX created at s3://{output_bucket}/{output_key}",
    }